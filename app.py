from flask import Flask, request, jsonify, send_file, render_template
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import io
import base64
from datetime import datetime
from docx2pdf import convert

app = Flask(__name__, template_folder='Templates')

# Workspace path
WORKSPACE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(WORKSPACE, "Templates", "Dynamic_Quotation_Template.docx")

# Base directory for generated documents (use writable /tmp/ folder on Vercel)
if os.environ.get('VERCEL'):
    GENERATED_DIR = os.path.join("/tmp", "Generated_Documents")
else:
    GENERATED_DIR = os.path.join(WORKSPACE, "Generated_Documents")

# --- HELPER FUNCTIONS FOR WORD STYLING ---
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_table_borders(table):
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border thickness
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')  # Professional grey border
        tblBorders.append(border)
    table._tbl.tblPr.append(tblBorders)

def set_paragraph_background(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def update_paragraph_text_preserving_style(paragraph, new_text):
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    # Preserve style from first run
    first_run = paragraph.runs[0]
    
    # Clear other runs to avoid duplication
    for run in paragraph.runs[1:]:
        run.text = ""
        
    # Split text by newline and add breaks properly
    lines = new_text.split('\n')
    first_run.text = lines[0]
    for line in lines[1:]:
        first_run.add_break()
        first_run.add_text(line)

def set_row_cell_widths(row):
    widths = [Inches(0.5), Inches(3.6), Inches(0.8), Inches(1.3)]
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

def replace_docx_image(docx_path, new_image_path, target_zip_path="word/media/image1.jpeg"):
    import tempfile
    import zipfile
    import shutil
    
    temp_dir = tempfile.mkdtemp()
    try:
        temp_docx = os.path.join(temp_dir, "temp.docx")
        shutil.copy2(docx_path, temp_docx)
        
        with zipfile.ZipFile(temp_docx, 'r') as yin:
            with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as yout:
                for item in yin.infolist():
                    if item.filename == target_zip_path:
                        yout.write(new_image_path, target_zip_path)
                    else:
                        yout.writestr(item, yin.read(item.filename))
    finally:
        shutil.rmtree(temp_dir)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    try:
        if request.is_json:
            data = request.json
        else:
            data_str = request.form.get("data")
            if not data_str:
                return jsonify({"error": "No data provided"}), 400
            import json
            data = json.loads(data_str)

        doc_type = data.get("doc_type", "Quotation")
        client_name = data.get("client_name", "").strip()
        client_address = data.get("client_address", "").strip()
        contact_no = data.get("contact_no", "").strip()
        date_val = data.get("date", datetime.today().strftime('%d/%m/%Y'))
        ref_no = data.get("ref_no", "").strip()
        subject = data.get("subject", "").strip()
        intro_text = data.get("intro_text", "").strip()
        city = data.get("city", "Abu Dhabi").strip()
        validity = data.get("validity", "10 days").strip()
        show_validity = data.get("show_validity", True)
        
        items = data.get("items", [])
        override_total = data.get("override_total")
        vat_enabled = data.get("vat_enabled", False)
        vat_rate = float(data.get("vat_rate", 5))
        
        terms = data.get("terms", "").strip()
        account_details = data.get("account_details", "").strip()
        footer_text = data.get("footer_text", "").strip()
        footer_color = data.get("footer_color", "0d05fa").replace("#", "")
        
        format_type = data.get("format", "pdf").lower()  # 'pdf' or 'docx'
        version = str(data.get("version", "2.0"))        # '1.0' (Original) or '2.0' (V2.0)

        # Extended Report, Photo & Contract Toggles
        photos = data.get("photos", [])
        show_pricing = data.get("show_pricing", True)
        show_photos = data.get("show_photos", True)
        show_scope = data.get("show_scope", True)
        show_contract = data.get("show_contract", False)
        show_advance = data.get("show_advance", False)
        show_terms = data.get("show_terms", True)
        show_account = data.get("show_account", True)

        advance_amount = float(data.get("advance_amount", 0))
        remaining_amount = float(data.get("remaining_amount", 0))

        scope_title = data.get("scope_title", "").strip()
        scope_text = data.get("scope_text", "").strip()

        if not client_name:
            return jsonify({"error": "Client Name is required"}), 400

        # Step 1: Render placeholders (logo & basic details)
        doc_tpl = DocxTemplate(TEMPLATE_PATH)
        context = {
            'client_name': client_name,
            'client_address': client_address,
            'date': date_val
        }
        doc_tpl.render(context)
        
        temp_dir = "/tmp" if os.environ.get('VERCEL') else WORKSPACE
        temp_filename = os.path.join(temp_dir, "temp_render.docx")
        doc_tpl.save(temp_filename)

        # Step 2: Open and apply edits to paragraphs
        doc_final = Document(temp_filename)

        # Find and replace headers/metadata
        heading_p = None
        for paragraph in doc_final.paragraphs:
            txt = paragraph.text.strip()
            
            # 1. Update UAE City
            if "The Manager" in txt and "UAE." in txt:
                update_paragraph_text_preserving_style(paragraph, f"The Manager {city}, UAE.")
                
            # 2. Update Date, Type, Ref
            elif txt.startswith("Date:"):
                new_text = f"Date: {date_val}   {doc_type}\nRef: {ref_no}" if ref_no else f"Date: {date_val}   {doc_type}"
                update_paragraph_text_preserving_style(paragraph, new_text)
                
            # 3. Update Validity
            elif txt.startswith("Validity:"):
                if show_validity and validity:
                    update_paragraph_text_preserving_style(paragraph, f"Validity: {validity}")
                else:
                    update_paragraph_text_preserving_style(paragraph, "")
                
            # 4. Update Main Heading (centered, search runs)
            elif txt in ["PAYMENT RECEIPT", "QUOTATION", "INVOICE"]:
                heading_p = paragraph
                for run in paragraph.runs:
                    cleaned = run.text.strip()
                    if cleaned in ["PAYMENT RECEIPT", "QUOTATION", "INVOICE"]:
                        run.text = run.text.replace(cleaned, doc_type.upper())
                
            # 5. Update Subject
            elif txt.startswith("Sub:"):
                update_paragraph_text_preserving_style(paragraph, f"Sub: {subject}")
                
            # 6. Update Intro Paragraph
            elif (txt.startswith("With reference to") or 
                  txt.startswith("We are pleased to acknowledge") or 
                  txt.startswith("This report confirms") or
                  txt.startswith("Following our") or
                  (len(txt) > 40 and "maintenance work" in txt.lower() and txt.startswith("Dear") is False)):
                update_paragraph_text_preserving_style(paragraph, intro_text)

        # Step 3: Insert Site Photos right under Document Heading Badge
        if show_photos and photos:
            num_photos = len(photos)
            cols = 4 if num_photos >= 4 else (2 if num_photos >= 2 else 1)
            rows = (num_photos + cols - 1) // cols
            
            photo_table = doc_final.add_table(rows=rows, cols=cols)
            photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for i, photo_obj in enumerate(photos):
                r_idx = i // cols
                c_idx = i % cols
                cell = photo_table.cell(r_idx, c_idx)
                
                data_url = photo_obj.get("dataUrl", "")
                caption = photo_obj.get("caption", "").strip()
                
                if data_url and "," in data_url:
                    try:
                        header, encoded = data_url.split(",", 1)
                        img_bytes = base64.b64decode(encoded)
                        img_stream = io.BytesIO(img_bytes)
                        
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        img_w = Inches(1.4 if cols >= 4 else (1.8 if cols == 2 else 2.0))
                        run.add_picture(img_stream, width=img_w)
                        
                        if caption:
                            p_cap = cell.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_cap = p_cap.add_run(caption)
                            run_cap.font.size = Pt(8.5)
                            run_cap.font.bold = True
                            run_cap.font.color.rgb = RGBColor(55, 65, 81)
                    except Exception as img_err:
                        print(f"Error processing image {i}: {img_err}")
            
            if heading_p:
                heading_p._p.addnext(photo_table._tbl)

        # Step 4: Insert Scope & Work Details into Word document
        if show_scope and (scope_title or scope_text):
            doc_final.add_paragraph("")
            if scope_title:
                p_stitle = doc_final.add_paragraph()
                run_stitle = p_stitle.add_run(scope_title)
                run_stitle.bold = True
                run_stitle.font.size = Pt(11.5)
                run_stitle.font.color.rgb = RGBColor(13, 5, 250)
            
            if scope_text:
                for line in scope_text.split('\n'):
                    line_str = line.strip()
                    if not line_str:
                        continue
                    if line_str.startswith("•") or line_str.startswith("-") or line_str.startswith("*"):
                        content = line_str[1:].strip()
                        p_bullet = doc_final.add_paragraph()
                        p_bullet.add_run("• ")
                        if ":" in content and not content.startswith("http"):
                            parts = content.split(":", 1)
                            r_b = p_bullet.add_run(parts[0] + ":")
                            r_b.bold = True
                            p_bullet.add_run(parts[1])
                        else:
                            p_bullet.add_run(content)
                    else:
                        p_norm = doc_final.add_paragraph()
                        if ":" in line_str and len(line_str) < 80:
                            parts = line_str.split(":", 1)
                            r_b = p_norm.add_run(parts[0] + ":")
                            r_b.bold = True
                            p_norm.add_run(parts[1])
                        else:
                            p_norm.add_run(line_str)

        # Step 5: Insert Line Items Pricing Table with Custom Row Shading & Lump Sum Override
        if show_pricing:
            doc_final.add_paragraph("")
            table = doc_final.add_table(rows=1, cols=4)
            set_table_borders(table)

            headers = ['No', 'Description', 'Qty', 'Total (AED)']
            hdr_cells = table.rows[0].cells
            set_row_cell_widths(table.rows[0])
            for i in range(4):
                hdr_cells[i].text = headers[i]
                set_cell_background(hdr_cells[i], 'D9D9D9')
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            calculated_subtotal = 0.0
            for idx, item in enumerate(items):
                no_str = f"{idx + 1:02d}"
                desc = item.get("desc", "").strip()
                qty = item.get("qty", "").strip()
                bg_style = item.get("bgStyle", "none")
                has_amount = item.get("has_amount", False)
                
                try:
                    amount = float(item.get("amount", 0))
                except (ValueError, TypeError):
                    amount = 0.0

                calculated_subtotal += amount
                    
                row_cells = table.add_row().cells
                set_row_cell_widths(table.rows[-1])
                row_cells[0].text = no_str
                row_cells[1].text = desc
                row_cells[2].text = qty
                
                # Display amount if provided/greater than 0, otherwise keep cell empty for lump sum items!
                if has_amount or amount > 0:
                    row_cells[3].text = f"{amount:.2f}"
                else:
                    row_cells[3].text = ""
                
                # Apply row background shading
                if bg_style == "grey":
                    for c in row_cells:
                        set_cell_background(c, 'D9D9D9')
                        if c.paragraphs and c.paragraphs[0].runs:
                            c.paragraphs[0].runs[0].font.bold = True
                            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                elif bg_style == "sub_grey":
                    for c in row_cells:
                        set_cell_background(c, 'F2F2F2')
                        if c.paragraphs and c.paragraphs[0].runs:
                            c.paragraphs[0].runs[0].font.bold = True
                            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                elif bg_style == "blue":
                    for c in row_cells:
                        set_cell_background(c, '0D05FA')
                        if c.paragraphs and c.paragraphs[0].runs:
                            c.paragraphs[0].runs[0].font.bold = True
                            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

            # Determine final subtotal (Override lump sum or calculated sum)
            if override_total is not None:
                try:
                    subtotal = float(override_total)
                except (ValueError, TypeError):
                    subtotal = calculated_subtotal
            else:
                subtotal = calculated_subtotal

            if vat_enabled:
                vat_amount = subtotal * (vat_rate / 100.0)
                grand_total = subtotal + vat_amount
                
                sub_cells = table.add_row().cells
                sub_cells[0].merge(sub_cells[2])
                sub_cells[0].text = "Subtotal Amount:"
                sub_cells[0].paragraphs[0].runs[0].font.bold = True
                sub_cells[3].text = f"{subtotal:.2f}"
                sub_cells[3].paragraphs[0].runs[0].font.bold = True
                set_cell_background(sub_cells[0], 'F2F2F2')
                set_cell_background(sub_cells[3], 'F2F2F2')
                
                vat_cells = table.add_row().cells
                vat_cells[0].merge(vat_cells[2])
                vat_cells[0].text = f"VAT ({vat_rate}%):"
                vat_cells[0].paragraphs[0].runs[0].font.bold = True
                vat_cells[3].text = f"{vat_amount:.2f}"
                vat_cells[3].paragraphs[0].runs[0].font.bold = True
                set_cell_background(vat_cells[0], 'F2F2F2')
                set_cell_background(vat_cells[3], 'F2F2F2')
            else:
                grand_total = subtotal

            tot_cells = table.add_row().cells
            tot_cells[0].merge(tot_cells[2])
            tot_cells[0].text = "Total Amount:" if not vat_enabled else "Grand Total:"
            tot_cells[0].paragraphs[0].runs[0].font.bold = True
            tot_cells[3].text = f"{grand_total:.2f}"
            tot_cells[3].paragraphs[0].runs[0].font.bold = True
            set_cell_background(tot_cells[0], 'F2F2F2')
            set_cell_background(tot_cells[3], 'F2F2F2')

            words_total = data.get("words_total", "").strip()
            words_cells = table.add_row().cells
            words_cells[0].merge(words_cells[3])
            words_cells[0].text = f"Amount in words AED: {words_total}"
            words_cells[0].paragraphs[0].runs[0].font.bold = True
            set_cell_background(words_cells[0], 'F2F2F2')

        # Step 6: Contract Clauses & Dual Signatures (if enabled)
        if show_contract:
            doc_final.add_paragraph("")
            
            # Notes or Special comments banner
            p_comm = doc_final.add_paragraph()
            r_comm = p_comm.add_run("Notes or Special comments")
            r_comm.bold = True
            r_comm.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(p_comm, '0D05FA')

            # Inclusion
            doc_final.add_paragraph("")
            p_inc = doc_final.add_paragraph()
            r_inc = p_inc.add_run("Inclusion :")
            r_inc.bold = True
            r_inc.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(p_inc, '0D05FA')

            inclusions = [
                "1. Checking of PH & Chlorine level and adjust as required supply of required chemicals.",
                "2. Check level of water in the pool, fill as required (water to be provided by client).",
                "3. Vacuum of the pool using vacuum head and telescopic handle and brush the walls as required.",
                "4. Check filtration system for pressure reading. Backwash the filters as required.",
                "5. Clean all debris from the pre-filter baskets as required.",
                "6. Check and ensure that all timers, underwater pool lights, all electrical controls & equipment's are in proper operation then report irregularities to client."
            ]
            for inc in inclusions:
                doc_final.add_paragraph(inc)

            # Exclusion
            doc_final.add_paragraph("")
            p_exc = doc_final.add_paragraph()
            r_exc = p_exc.add_run("Exclusion :")
            r_exc.bold = True
            r_exc.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(p_exc, '0D05FA')

            exclusions = [
                "1. Repairs/ rectification works such as pool leakage, re-grouting of tiles, replacement of underwater lights, filters and Pumps Repairs/rectification and spare parts will be charged on separate cost and subject to clients approval before proceeding with the required works.",
                "2. Snag and all materials",
                "3. All other works not related to pool cleaning and checking of equipment"
            ]
            for exc in exclusions:
                doc_final.add_paragraph(exc)

            # Notes
            doc_final.add_paragraph("")
            p_notes = doc_final.add_paragraph()
            r_notes = p_notes.add_run("Notes")
            r_notes.bold = True
            r_notes.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(p_notes, '0D05FA')

            notes_list = [
                "1. A separate quotation for items under exclusions (if required) will be provided to client for approval",
                "2. Services provided Monday, Thursday, Friday, excluding public holidays.",
                "3. Client is required to give access to the pool",
                "4. Our company is not liable for any direct or consequential costs or liabilities arising due to failure or malfunction of any part of the pool or pool equipment."
            ]
            for note in notes_list:
                doc_final.add_paragraph(note)

            # Advance 100% badge
            doc_final.add_paragraph("")
            p_adv = doc_final.add_paragraph()
            r_adv = p_adv.add_run("Payment Advance 100%")
            r_adv.bold = True
            r_adv.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(p_adv, '0D05FA')

            # Dual Signature Block (2 columns)
            doc_final.add_paragraph("")
            sig_table = doc_final.add_table(rows=1, cols=2)
            sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            c_left = sig_table.cell(0, 0)
            c_right = sig_table.cell(0, 1)
            c_left.width = Inches(3.5)
            c_right.width = Inches(3.0)
            
            p_l = c_left.paragraphs[0]
            p_l.add_run("Sincerely,\nIjaz Hussian\nFit Pool General Building Maintenance LLC\nAbu Dhabi, Dubai\n\n\nSigned for and on behalf of Customer\nName and Signature:\nDate:")
            
            p_r = c_right.paragraphs[0]
            p_r.add_run("Account details\nIjaz Hussain\nMeshruq Bank\nIBAN: AE 660330000019200061112\nContact No: +971564378296")

        # Step 7: Advance & Remaining Payment Breakdown
        if show_advance:
            doc_final.add_paragraph("")
            adv_p = doc_final.add_paragraph()
            r_adv1 = adv_p.add_run("Advance Payment Received: ")
            r_adv1.bold = True
            adv_p.add_run(f"{advance_amount:.2f} AED\n")
            r_adv2 = adv_p.add_run("Remaining Payment Balance: ")
            r_adv2.bold = True
            adv_p.add_run(f"{remaining_amount:.2f} AED")
            set_paragraph_background(adv_p, 'F8FAFC')

        # Step 8: Append Terms & Conditions (Standard)
        if show_terms and terms:
            term_p = doc_final.add_paragraph()
            term_p.paragraph_format.space_before = Pt(6)
            term_p.paragraph_format.space_after = Pt(2)
            term_p.add_run("Payment term and condition").bold = True
            
            p_t_content = doc_final.add_paragraph(terms)
            p_t_content.paragraph_format.space_before = Pt(0)
            p_t_content.paragraph_format.space_after = Pt(4)

        # Step 9: Append Bank Account Details (Standard)
        if not show_contract and show_account and account_details:
            acc_p = doc_final.add_paragraph()
            acc_p.paragraph_format.space_before = Pt(6)
            acc_p.paragraph_format.space_after = Pt(2)
            acc_p.add_run("Account details").bold = True
            
            p_a_content = doc_final.add_paragraph(account_details)
            p_a_content.paragraph_format.space_before = Pt(0)
            p_a_content.paragraph_format.space_after = Pt(4)

        # Step 10: Add Footer Banner
        if footer_text:
            footer_para = doc_final.add_paragraph()
            footer_para.paragraph_format.space_before = Pt(8)
            footer_para.paragraph_format.space_after = Pt(0)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_run = footer_para.add_run(footer_text)
            footer_run.bold = True
            footer_run.font.color.rgb = RGBColor(255, 255, 255)
            set_paragraph_background(footer_para, footer_color)

        # Determine folders and filenames
        safe_name = "".join(c if c.isalnum() or c in [' ', '_', '-'] else "" for c in client_name).replace(" ", "_")
        date_str = datetime.today().strftime('%d-%m-%Y')
        
        doc_type_lower = doc_type.lower()
        if "invoice" in doc_type_lower:
            folder = os.path.join(GENERATED_DIR, "Invoices")
            prefix = "Invoice"
        elif "receipt" in doc_type_lower:
            folder = os.path.join(GENERATED_DIR, "Receipts")
            prefix = "Payment_Receipt"
        elif "contract" in doc_type_lower:
            folder = os.path.join(GENERATED_DIR, "Contracts")
            prefix = "Cleaning_Contract"
        elif "report" in doc_type_lower or "scope" in doc_type_lower:
            folder = os.path.join(GENERATED_DIR, "Reports")
            prefix = doc_type.replace(" ", "_")
        else:
            folder = os.path.join(GENERATED_DIR, "Quotations")
            prefix = "Quotation"
            
        os.makedirs(folder, exist_ok=True)
        
        docx_path = os.path.join(folder, f"{prefix}_{safe_name}_{date_str}.docx")
        pdf_path = os.path.join(folder, f"{prefix}_{safe_name}_{date_str}.pdf")
        
        doc_final.save(docx_path)
        
        # Swapping the background image for Version 2.0
        if version == "2.0":
            v2_image_path = os.path.join(WORKSPACE, "static", "image1_v2.jpeg")
            if os.path.exists(v2_image_path):
                replace_docx_image(docx_path, v2_image_path)
        
        # Cleanup temporary render
        if os.path.exists(temp_filename):
            os.remove(temp_filename)

        # If PDF is requested
        if format_type == "pdf":
            if os.name == 'nt':
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    convert(docx_path, pdf_path)
                finally:
                    pythoncom.CoUninitialize()
            else:
                try:
                    convert(docx_path, pdf_path)
                except Exception as e:
                    raise RuntimeError(
                        "PDF conversion failed. This is likely because LibreOffice/MS Word is not installed on the server (Vercel). "
                        "Please download the document in DOCX format instead."
                    ) from e

        target_path = pdf_path if format_type == "pdf" else docx_path
        download_filename = os.path.basename(target_path)

        response = send_file(target_path, as_attachment=True, download_name=download_filename)
        response.headers["Content-Disposition"] = f'attachment; filename="{download_filename}"'
        response.headers["Access-Control-Expose-Headers"] = "Content-Disposition"
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route('/download/<folder_name>/<filename>')
def download_file(folder_name, filename):
    try:
        if folder_name not in ['Quotations', 'Invoices', 'Receipts', 'Reports', 'Contracts']:
            return jsonify({"error": "Invalid folder"}), 400
        
        safe_filename = os.path.basename(filename)
        file_path = os.path.abspath(os.path.join(GENERATED_DIR, folder_name, safe_filename))
        
        expected_dir = os.path.abspath(os.path.join(GENERATED_DIR, folder_name))
        if not file_path.startswith(expected_dir):
            return jsonify({"error": "Access denied"}), 403
            
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
            
        response = send_file(file_path, as_attachment=True, download_name=safe_filename)
        response.headers["Content-Disposition"] = f'attachment; filename="{safe_filename}"'
        return response
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
