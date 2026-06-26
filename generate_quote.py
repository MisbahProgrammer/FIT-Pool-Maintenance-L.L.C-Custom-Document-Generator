from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

# --- MAGIC FUNCTION TO COLOR TABLE CELLS ---
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# --- MAGIC FUNCTION TO DRAW TABLE BORDERS ---
def set_table_borders(table):
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # Border thickness
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # Black color
        tblBorders.append(border)
    table._tbl.tblPr.append(tblBorders)

# --- MAGIC FUNCTION TO COLOR PARAGRAPH BACKGROUND ---
def set_paragraph_background(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def create_quotation():
    print("=== FIT POOL QUOTATION GENERATOR ===")
    client_name = input("Enter Client/Company Name: ")
    client_address = input("Enter Address/Apt Details: ")

    items = []
    grand_total = 0
    item_counter = 1

    print("\n--- Enter items (Type 'done' in the description to finish) ---")
    while True:
        desc = input(f"Item {item_counter} Description: ")
        if desc.lower() == "done":
            break

        qty = input(f"Item {item_counter} Qty (e.g 2 No's): ")

        while True:
            try:
                amount = float(input(f"Item {item_counter} Amount (AED): "))
                break
            except ValueError:
                print("Please enter a valid number for the amount.")

        items.append({
            'no': f"{item_counter:02d}",
            'desc': desc,
            'qty': qty,
            'amount': amount
        })

        grand_total += amount
        item_counter += 1

    print(f"\nCalculated grand total: {grand_total} AED")
    words_total = input("Please type the grand total in words (e.g Three Thousand...): ")

    safe_name = client_name.replace(" ", "_")
    final_filename = f"Quotation_{safe_name}_{datetime.today().strftime('%d-%m-%Y')}.docx"
    temp_filename = "temp_render.docx"

    # ==========================================
    # PART 1: Fill in the Top Half (Logo & Client Info)
    # ==========================================
    doc_tpl = DocxTemplate("Dynamic_Quotation_Template.docx")
    context = {
        'client_name': client_name,
        'client_address': client_address,
        'date': datetime.today().strftime('%d/%m/%Y')
    }
    doc_tpl.render(context)
    doc_tpl.save(temp_filename)

    # ==========================================
    # PART 2: Draw the Table and Footer
    # ==========================================
    doc_final = Document(temp_filename)
    
    # Add a little spacing before table
    doc_final.add_paragraph("")

    # 1. Create the Table and Force Borders
    table = doc_final.add_table(rows=1, cols=4)
    set_table_borders(table)

    # 2. Style the Header Row (Grey background, Bold Black text)
    headers = ['No', 'Description', 'Qty', 'Total Amount /AED']
    hdr_cells = table.rows[0].cells
    
    for i in range(4):
        hdr_cells[i].text = headers[i]
        set_cell_background(hdr_cells[i], 'D9D9D9') # Professional Light Grey
        
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0) # Black text

    # 3. Add the Item Rows Dynamically
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item['no']
        row_cells[1].text = item['desc']
        row_cells[2].text = item['qty']
        row_cells[3].text = f"{item['amount']:.2f}"

    # 4. Add the Totals INSIDE the table (Merged Cells)
    # Row for Total Amount
    total_row = table.add_row().cells
    total_row[0].merge(total_row[2])
    total_row[0].text = "Total Amount:"
    total_row[0].paragraphs[0].runs[0].font.bold = True
    total_row[3].text = f"{grand_total:.2f}"
    total_row[3].paragraphs[0].runs[0].font.bold = True
    
    set_cell_background(total_row[0], 'F2F2F2') 
    set_cell_background(total_row[3], 'F2F2F2')

    # Row for Amount in Words
    words_row = table.add_row().cells
    words_row[0].merge(words_row[3])
    words_row[0].text = f"Amount in words AED: {words_total}"
    words_row[0].paragraphs[0].runs[0].font.bold = True
    set_cell_background(words_row[0], 'F2F2F2')

    # 5. Add the Footer Information below the table
    doc_final.add_paragraph("") # Blank space
    
    term_p = doc_final.add_paragraph()
    term_p.add_run("Payment term and condition").bold = True
    doc_final.add_paragraph("75% Advance payment\n25% Payment after completion of work")

    doc_final.add_paragraph("") # Blank space

    acc_p = doc_final.add_paragraph()
    acc_p.add_run("Account details").bold = True
    doc_final.add_paragraph("Ijaz Hussain\nMeshruq Bank\nIBAN: AE 660330000019200061112\nContact No: +971564378296")
    
    doc_final.add_paragraph("") # Blank space before banner

    # --- THE NEW COLORED FOOTER BANNER ---
    footer_para = doc_final.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center the text
    
    footer_run = footer_para.add_run("Thanks and Regards: Fit Pool Building Maintenance L.L.C")
    footer_run.bold = True
    footer_run.font.color.rgb = RGBColor(255, 255, 255) # White text
    
    set_paragraph_background(footer_para, '0d05fa')
    # ==========================================
    # PART 3: Save and Clean up
    # ==========================================
    doc_final.save(final_filename)

    if os.path.exists(temp_filename):
        os.remove(temp_filename)

    print(f"\nSUCCESS! Your fully automated quotation is ready: {final_filename}")

if __name__ == "__main__":
    create_quotation()